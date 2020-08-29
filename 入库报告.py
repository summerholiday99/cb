# -*- coding: utf-8 -*-

import datetime as dt
import pandas as pd
import warnings
warnings.filterwarnings("ignore")
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from functools import reduce

def produce_file(str_raw):

    dfs = pd.read_excel(str_road, sheet_name=None)
    temp = dfs['简介']
    srt_filename = '{}  {}'.format(temp.loc['转债代码','value'],
                                   temp.loc['转债简称','value'])
    str_out = r'{}\可转债\转债入库\{}.docx'.format(str_raw,srt_filename)
     
    document = docx.Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 文件头
    p = document.add_paragraph(srt_filename)
    p = document.add_paragraph('公司：{}'.format(temp.loc['正股名称','value']))
    p = document.add_paragraph('债券规模：{}亿'.format(temp.loc['转债余额','value']))
    p = document.add_paragraph('期限：{}年'.format(temp.loc['发行期限','value'],
                                                temp.loc['转债余额','value']))
    p = document.add_paragraph('外部评级：{}  {}/{} '.format(
                               temp.loc['评级机构','value'],
                               temp.loc['主体评级','value'],
                               temp.loc['信用等级','value']))

    # 主体部分
    p = document.add_paragraph('')
    p.add_run('公司信用资质简要分析').bold = True
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph(_st_introduction(dfs))      #公司简介 股权结构
    p.paragraph_format.first_line_indent = Inches(0.3)
    p = document.add_paragraph('做什么生意？行业空间及增速 产业链（上游及下游） 竞争格局')
    p.paragraph_format.first_line_indent = Inches(0.3)
    p = document.add_paragraph(_st_financial_summary(dfs)+_st_main_business(dfs))
    p.paragraph_format.first_line_indent = Inches(0.3)     # 财务摘要+主营构成
    p = document.add_paragraph(_cb_info(dfs))              # 转债规模、条款、募投项目
    p.paragraph_format.first_line_indent = Inches(0.3)
    p = document.add_paragraph(_st_financial_risk(dfs))    #'资产负债表&信用风险'
    p.paragraph_format.first_line_indent = Inches(0.3)
    p = document.add_paragraph('综上，发行人信用风险较低。')
    p.paragraph_format.first_line_indent = Inches(0.3)

    # 承诺函
    p = document.add_paragraph('')
    p.add_run('承诺函').bold = True
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(srt_filename+'研究报告是本人根据公开市场的各类信息，'+
                               '在充分研究论证的基础上得出的结论，本人承诺在此次研究过程中'+
                               '不存在利用内幕信息及未公开信息的情况。对于因履行工作职责所'+
                               '知悉的内幕信息、未公开信息，本人承诺将根据公司制度的规定履行'+
                               '报告、知情人登记及保密义务，防止内幕信息、未公开信息的进一步'+
                               '不当传播和使用。')
    p.paragraph_format.first_line_indent = Inches(0.3)
    p = document.add_paragraph('研究员：黄基力')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = document.add_paragraph('时间：{}'.format(str(dt.datetime.today().date())))
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.save(str_out)
    return

def _st_introduction(dfs):
    '''
    这一段写简介: 公司、股权结构、持有人、质押、解禁
    :param dfs:
    :return:
    '''

    temp = dfs['简介']
    temp2 = dfs['股权&股权风险']

    s1 = '{}属于{}行业，主营产品{}。'.format(temp.loc['正股名称','value'],
                                   temp.loc['申万行业','value'],
                                   temp.loc['主营产品类型', 'value'])
    s2 = '公司实际控制人为{}（{}），' \
         '持股比例{}。公司董事长为{}。'.format(temp2.loc['实际控制人', 'value'],
                                  temp2.loc['实际控制人属性', 'value'],
                                  temp2.loc['实控人持股比例', 'value'],
                                  temp2.loc['董事长', 'value'], )

    s3 = '公司前三大股东分别为{}、{}、{}，' \
         '持股比例分别为{:.2f}%、{:.2f}%、{:.2f}%。'.format(temp2.loc['第一大股东', 'value'],
                                    temp2.loc['第二大股东', 'value'],
                                    temp2.loc['第三大股东', 'value'],
                                    temp2.loc['第一大股东持股比例', 'value'],
                                    temp2.loc['第二大股东持股比例', 'value'],
                                    temp2.loc['第三大股东持股比例', 'value'])

    s3_1 = '基金持股比例合计为{:.2f}%。'.format(temp2.loc['基金持股比例合计', 'value'])

    s4 = '质押方面，前三大股东及实控人累计质押占持股比例分别为' \
         '{:.2f}%、{:.2f}%、{:.2f}%，质押风险？。'.format(temp2.loc['第一大股东累计质押占持股比例', 'value'],
                                                temp2.loc['第二大股东累计质押占持股比例', 'value'],
                                                temp2.loc['第三大股东累计质押占持股比例', 'value'],
                                                temp2.loc['实控人累计质押占持股比例', 'value'],)
    s5 = '解禁方面，未来最近一次解禁为{}，{}解禁，' \
         '占流动股比例{:.2%}，解禁压力？。'.format(temp2.loc['指定日之后最近一次解禁股份性质', 'value'],
                                  temp2.loc['指定日之后最近一次解禁日期', 'value'],
                                  temp2.loc['指定日之后最近一次解禁数量占流动股比例', 'value'])

    list_str = [s1, s2, s3, s3_1, s4, s5]
    str_sentence = reduce(lambda x, y: x + y, list_str)
    return str_sentence

def _st_financial_summary(dfs):
    '''
    源excel数据做成标准格式比较难，这边就靠定位了
    :param dfs:
    :return:
    '''
    def f(dd):
        return str(dd.year) + '/' +str(dd.month) if dd.month==6 else str(dd.year)
    temp = dfs['财务摘要']
    str_year = f(temp.loc[3, 'Unnamed: 4']) + '-' + f(temp.loc[3, 'Unnamed: 6'])

    # 营收及同比
    s1 = '{},公司分别实现营收{}亿' \
         '（YOY {:.2f}%）、{}亿（YOY {:.2f}%）、{}亿（YOY {:.2f}%）;'.format(str_year,
                                        temp.loc[7, 'Unnamed: 4'],temp.loc[8, 'Unnamed: 4'],
                                        temp.loc[7, 'Unnamed: 5'], temp.loc[8, 'Unnamed: 5'],
                                        temp.loc[7, 'Unnamed: 6'], temp.loc[8, 'Unnamed: 6'],
                                        )
    # 归母净利及同比
    s2 = '实现归母净利{}亿（YOY {:.2f}%）、' \
         '{}亿（YOY {:.2f}%）、{}亿（YOY {:.2f}%）;'.format(temp.loc[15, 'Unnamed: 4'], temp.loc[16, 'Unnamed: 4'],
                                   temp.loc[15, 'Unnamed: 5'], temp.loc[16, 'Unnamed: 5'],
                                   temp.loc[15, 'Unnamed: 6'], temp.loc[16, 'Unnamed: 6'])
    # 毛利率
    s3 = '毛利率分别为{:.2f}%、{:.2f}%、{:.2f}%；'.format(temp.loc[57, 'Unnamed: 4'],
                                              temp.loc[57, 'Unnamed: 5'],
                                              temp.loc[57, 'Unnamed: 6'],)
    # 净利率
    s4= '净利率为{:.2f}%、{:.2f}%、{:.2f}%；'.format(temp.loc[58, 'Unnamed: 4'],
                                              temp.loc[58, 'Unnamed: 5'],
                                              temp.loc[58, 'Unnamed: 6'],)
    # ROE
    s5 = 'ROE为{:.2f}%、{:.2f}%、{:.2f}%。'.format(temp.loc[52, 'Unnamed: 4'],
                                               temp.loc[52, 'Unnamed: 5'],
                                               temp.loc[52, 'Unnamed: 6'], )
    list_str = [s1, s2, s3, s4, s5]
    str_sentence = reduce(lambda x, y: x + y, list_str)
    return str_sentence

def _st_main_business(dfs):
    '''
    分版块收入，及毛利率(位置原因，有一些筛选计算)
    这个wind分类口径可能不太对，不对你就自己做
    同一名字一般四个值：分别是营收 成本 毛利 毛利率
    :param dfs:
    :return:
    '''
    def pick(temp,row_name,values_rank):
        return  temp.loc[temp['证券代码']==row_name,
                         str_name].values[values_rank]

    temp: pd.DataFrame = dfs['主营构成']
    str_name = 'Unnamed: {}'.format(str(temp.shape[1]-2)
                                    if (temp.loc[temp['证券代码']=='        报告期',
                                                'Unnamed: 6'] == '中报').values[0]
                                    else str(temp.shape[1]-1))                 # 用最新的全年数据,不用半年（有的半年不全）
    list_scale = [temp.loc[temp['证券代码']=='                产品',:].index[0],
                  temp.loc[temp['证券代码']=='                地区',:].index[0]]  # 主营业务可能存在的位置
    list_loc = temp.loc[list(range(list_scale[0]+1, list_scale[1])),            # 主营业务存在的位置，这样就有了名字
                        str_name].dropna().index.values
    list_name = temp.loc[list_loc,'证券代码'].values                             # 主营业务名称

    dict_word = {ii.strip():[pick(temp, ii, 0),pick(temp, ii, 3)]
                 for ii in list_name}

    s1 = '从最新年报看，公司总营收{}亿，其中'.format(pick(temp, '        营业总收入', 0)) #总营收
    s2 = ['{}实现营收{}亿，毛利率{:.2f}%；'.format(ii,dict_word[ii][0],dict_word[ii][1])  for ii in dict_word.keys()]
    s2[-1] = s2[-1].replace('；','。')
    s2 = reduce(lambda x, y: x + y, s2)
    str_sentence = s1 + s2
    return str_sentence

def _cb_info(dfs):

    temp = dfs['简介']
    temp2 = dfs['转债条款']

    s1 = '此次发行转债，等级{}，期限{}年，发行规模为{}亿，' \
         '稀释比率{:.2f}%。'.format(temp.loc['信用等级','value'],
                        temp.loc['发行期限', 'value'],
                        temp.loc['转债余额', 'value'],
                        temp2.loc['稀释比率', 'value'],)

    s2 = '条款上，转股起始时间{}，转股价格{}元' \
         '（当前正股价格{}元）；'.format(temp2.loc['自愿转股起始时间', 'value'],
                              temp2.loc['转股价格', 'value'],
                              temp2.loc['正股价格', 'value'],)

    s3 = '赎回条款为{}，{}/{}；'.format(temp2.loc['赎回触发比率', 'value'],
                                temp2.loc['赎回触发计算时间区间', 'value'],
                                temp2.loc['赎回触发计算最大时间区间', 'value'],)

    s4 = '回售起始时间{}，条款为{}，{}/{}；'.format(temp2.loc['条件回售起始时间', 'value'],
                                       temp2.loc['回售触发比率', 'value'],
                                       temp2.loc['回售触发计算时间区间', 'value'],
                                       temp2.loc['回售触发计算最大时间区间', 'value'],)

    s5 = '利率条款为{}利率补偿为{}'.format(temp2.loc['利率条款', 'value'],
                             temp2.loc['利率补偿条款', 'value'],)
    s6 = '募投项目上，此次募集{}亿，其中？亿投向？项目，其中？亿投向？项目。'.format(temp.loc['转债余额', 'value'])

    list_str = [s1, s2, s3, s4, s5, s6]
    str_sentence = reduce(lambda x, y: x + y, list_str)
    return str_sentence

def _st_financial_risk(dfs):

    def f(dd):
        return str(dd.year) + '/' +str(dd.month) if dd.month==6 else str(dd.year)

    temp = dfs['三表']
    temp2 = dfs['财务风险'].set_index('报告参数')
    col_t = temp2.columns[-1]
    col_t_s = temp2.columns[-3:]
    str_t = '{}-{}'.format(f(temp2.columns[-3]), f(temp2.columns[-1]))

    s1 = '{},公司总资产{:.2f}亿，流动资产中，货币{:.2f}亿，' \
         '应收账款+票据{:.2f}亿，存货{:.2f}亿，' \
         '其他流动资产、长期应收款、长期股权投资分别为{:.2f}亿、{:.2f}亿、{:.2f}亿；' \
         '非流动资产中，固定+在建{:.2f}亿，无形资产{:.2f}亿，商誉{:.2f}亿，' \
         '其他非流动资产{:.2f}亿。'.format(
                                        f(temp2.columns[-1]),
                                        temp.loc['        资产总计',col_t],
                                        temp.loc['        货币资金',col_t],
                                        temp.loc['        应收票据',col_t]+temp.loc['        应收账款',col_t],
                                        temp.loc['        预付款项',col_t],
                                        temp.loc['        存货',col_t],
                                        temp.loc['        其他流动资产',col_t],
                                        temp.loc['        长期应收款',col_t],
                                        temp.loc['        长期股权投资',col_t],
                                        temp.loc['        固定资产',col_t]+temp.loc['        在建工程',col_t],
                                        temp.loc['        无形资产',col_t],
                                        temp.loc['        商誉',col_t],
                                        temp.loc['        其他非流动资产',col_t],
                                        )

    s2_1 = '负债端，总债务{:.2f}亿，其中短期债务{:.2f}亿，' \
         '长期债务{:.2f}亿，资产负债率{:.2%}'.format(temp2.loc['总债务',col_t],
                                          temp2.loc['短期债务', col_t],
                                          temp2.loc['长期债务', col_t],
                                          temp2.loc['资产负债率', col_t],)

    list_1 = temp2.loc['总债务/(总债务+所有者权益）', col_t_s].values
    list_2 = temp2.loc['EBITDA/营业收入', col_t_s].values
    list_3 = temp2.loc['CFO(亿元）', col_t_s].values

    s2_2 = '{},全部债务资本化率分别为{:.2%}、{:.2%}、{:.2%}；' \
         'EBITDA/营业收入分别为{:.2%}、{:.2%}、{:.2%}，盈利能力？；' \
         '经营性现金流分别为{:.2%}、{:.2%}、{:.2%}，现金流状况？。'.format(str_t,
                                     list_1[0],list_1[1],list_1[2],
                                     list_2[0], list_2[1], list_2[2],
                                     list_3[0], list_3[1], list_3[2],
                                     )
    s2_3 = '{}包括货币、交易性金融资产及应收票据在内的现金资产{:.2f}亿,' \
         '短期债务{:.2f}亿，短期流动性压力？。'.format(f(temp2.columns[-1]),
                                        temp2.loc['现金类资产', col_t],
                                        temp2.loc['短期债务', col_t],
                                        )
    list_str = [s1, s2_1, s2_2, s2_3]
    str_sentence = reduce(lambda x, y: x + y, list_str)
    return str_sentence

if '__name__' == '__main__':
    str_raw = r'C:\Users\huangjili\Desktop\同步\我的坚果云'
    str_road = r'{}\cb\转债入库报告.xlsx'.format(str_raw)
    produce_file(str_raw)